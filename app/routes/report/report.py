from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from app.routes.dataview.gen_data import gen_month_table, gen_base_table
from reportlab.lib.pagesizes import landscape, A4
from flask_login import login_required
from flask import Blueprint, send_file, current_app
import app.global_vars as global_vars
from app.models import Center, Month
from reportlab.lib import colors
from docx import Document
from io import BytesIO
import os

report_bp = Blueprint(
                      'report',
                      __name__,
                      template_folder='templates',
                      static_folder='static',
                      static_url_path='/static/report'
                      )


@report_bp.route('/report/<center>/<month>/<year>')
@login_required
def gen_report(center, month, year):
    try:
        # Determine the path to the DOCX template
        template_path = os.path.join(current_app.root_path, 'static', 'escala.docx')
        
        # Log the path being used
        current_app.logger.info(f"Using template path: {template_path}")
        
        # Load the DOCX template
        doc = Document(template_path)

        # Replace placeholders in the document
        for paragraph in doc.paragraphs:
            if '{names}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{names}', "hello\nworld!\n!\nIam\na\nnew\nline")
            if '{center}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{center}', center)
            if '{month}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{month}', month)
            if '{year}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{year}', year)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '{names}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{names}', "hello\nworld!\n!\nIam\na\nnew\nline")
                        if '{center}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{center}', center)
                        if '{month}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{month}', month)
                        if '{year}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{year}', year)

        # Save the updated document to a BytesIO object
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Send the file as an attachment
        return send_file(file_stream, 
                         download_name='report.docx', 
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                         as_attachment=True)
    except Exception as e:
        current_app.logger.error(f"Error generating report: {e}")
        return str(e), 500



@report_bp.route('/print-table/<center_abbr>/<month_name>/<year>')
@login_required
def print_table(center_abbr, month_name, year):
    if month_name == 'null':
        center = Center.query.filter_by(abbreviation=center_abbr).first()
        if center is None:
            raise Exception(f"{center_abbr} not found")
        
        month = Month.get_current()

        data_table = gen_base_table(center_abbr=center_abbr, names_only=True, abbr_names=True)
    else:    
        month_num = global_vars.MESES.index(month_name) + 1
        month = Month.query.filter_by(number=month_num).first()
        if month is None:
            raise Exception(f"{month_name} not found")
        
        center = Center.query.filter_by(abbreviation=center_abbr).first()
        if center is None:
            raise Exception(f"{center_abbr} not found")
        
        data_table = gen_month_table(center_abbr, month_name, year, names_only=True)

    for i in range(len(data_table[0])):
        if i == 0:
            continue
        data_table[0][i] = data_table[0][i][0]

    first_header = ["HOSPITAL UNIVERSITÁRIO EVANGÉLICO MACKENZIE"] + [''] * (len(data_table[0]) - 1)
    
    second_header = ["Grupo de Anestesia Mackenzie"] + [''] * (len(data_table[0]) - 1)
    second_header[len(data_table[0])//3] = f"{center.name} - {center.abbreviation}"
    second_header[2*(len(data_table[0])//3)] = f"COMPETÊNCIA: {month.name}/{month.year}"

    third_header = [f"RESPONSÁVEL: {month.leader}"] + [''] * (len(data_table[0]) - 1)
    third_header[len(data_table[0])//2] = f"CURITIBA, {global_vars.STR_DAY}/{month.number}/{month.year}"

    data = [first_header, second_header] + data_table

    buffer = BytesIO()

    # Set the page size to landscape
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), topMargin=10)

    table = Table(data)
    style = TableStyle([
        ('FONT', (0, 0), (-1, -1), 'Helvetica', 8),  # Apply font size 8 to the entire table
        ('GRID', (0, 0), (-1, -1), 1, colors.black),  # Add grid (borders) to the table

        ('SPAN', (0, 0), (len(data[0])-1, 0)),

        ('SPAN', (0, 1), (len(data[0])//3 - 1, 1)),
        ('SPAN', (len(data[0])//3, 1), (2*len(data[0])//3 - 1, 1)),
        ('SPAN', (2*len(data[0])//3, 1), (-1, 1)),

        # ('SPAN', (0, 2), (len(data[0])//2 - 1, 2)),
        # ('SPAN', (len(data[0])//2, 2), (-1, 2)),

        # Aligning text in the merged cell
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
        ('ALIGN', (0, 1), (-1, 1), 'CENTER'),
        ('VALIGN', (0, 1), (-1, 1), 'MIDDLE'),
        # ('ALIGN', (0, 2), (-1, 2), 'CENTER'),
        # ('VALIGN', (0, 2), (-1, 2), 'MIDDLE'),
    ])

    # for col_index, cell_value in enumerate(data[1]):
    #     print(cell_value)
    # if cell_value == "":
    #   style.add('BACKGROUND', (col_index, 0), (col_index, -1), colors.lightgrey)

    table.setStyle(style)
    table.repeatRows = 4
    doc.build([table])
    buffer.seek(0)

    return send_file(buffer, download_name='report.pdf', mimetype='application/pdf', as_attachment=False)
